# -*- coding: UTF-8 -*-
'''
@Author  ：程序员晚枫，B站/抖音/微博/小红书/公众号
@WeChat     ：CoderWanFeng
@Blog      ：www.python-office.com
@Date    ：2022/12/30 10:57 
@Description     ：
'''
from docx import Document
# 飞件涉及显落样式修改
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# 文件造及文字样式修改，领色修改、字号调整
from docx.shared import RGBColor, Pt
# 设置中文字体
from docx.oxml.ns import qn
from openpyxl import load_workbook
from potime import RunTime


@RunTime
def request_for_leave_wanfeng(name, department, reason, days, date):
    doc = Document()
    heading_1 = '请假条'
    paragraph_1 = doc.add_heading(heading_1, level=1)  # 居中对芳
    paragraph_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 标题要打，单独修改教大字号
    for run in paragraph_1.runs:
        run.font.size = Pt(17)
    word_1 = "    本人"
    word_2 = "，所在部门"
    word_3 = "，由于"
    word_4 = "，需请假"
    word_5 = "天。"

    paragraph_3_python_office = doc.add_paragraph()
    paragraph_3_python_office.add_run(word_1)
    paragraph_3_python_office.add_run(name).underline = True
    paragraph_3_python_office.add_run(word_2)
    paragraph_3_python_office.add_run(department).underline = True
    paragraph_3_python_office.add_run(word_3)
    paragraph_3_python_office.add_run(reason).underline = True
    paragraph_3_python_office.add_run(word_4)
    paragraph_3_python_office.add_run(str(days)).underline = True
    paragraph_3_python_office.add_run(word_5)
    # 设置下划线
    paragraph_3_python_office.paragraph_format.line_spacing = 1.5
    word_6 = '申请人：'
    paragraph_4_python4office_cn = doc.add_paragraph()
    paragraph_4_python4office_cn.add_run(word_6)
    paragraph_4_python4office_cn.add_run(name).underline = True
    paragraph_4_python4office_cn.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    word_7 = '日期：'
    sign_date = "{}年{}月{}日".format(date.split('-')[0], date.split('-')[1], date.split('-')[2])
    paragraph_5_python_office_com = doc.add_paragraph()
    paragraph_5_python_office_com.add_run(word_7)
    paragraph_5_python_office_com.add_run(sign_date).underline = True
    paragraph_5_python_office_com.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # 统一修改颜色
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = '楷体'
            r = run._element.rPr.rFonts
            r.set(qn('w:eastAsia'), '楷体')

    doc.save(r".\res\{}-请假条.docx".format(name))


@RunTime
def read_excel_python_office():
    path = r'晚枫的Excel员工文件.xlsx'  # 路径为Excel 北件所在的位置，可按实际待况更改
    workbook = load_workbook(path)
    sheet = workbook.active
    n = 0
    for row in sheet.rows:
        if n:
            name = row[0].value
            department = row[1].value
            reason = row[2].value
            days = row[3].value
            date = str(row[4].value).split()[0]
            print(date)
            if date != 'None':
                request_for_leave_wanfeng(name, department, reason, days, date)
        n += 1


if __name__ == '__main__':
    read_excel_python_office()
